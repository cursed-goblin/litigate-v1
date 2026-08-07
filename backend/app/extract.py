from __future__ import annotations

import io
import re

try:
    import fitz
except ImportError:  # PyMuPDF exposes both module names
    import pymupdf as fitz

from docx import Document

_HYPHEN_BREAK = re.compile(r"(\w)-\n(\w)")
_PAGE_MARKER = re.compile(r"^\s*(?:page\s+)?\d{1,4}\s*(?:of\s+\d{1,4})?\s*$", re.IGNORECASE)
_TRAILING_SPACE = re.compile(r"[ \t]+$", re.MULTILINE)
_EXTRA_BLANKS = re.compile(r"\n{3,}")


class ExtractionError(RuntimeError):
    pass


def normalise(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ")
    text = _HYPHEN_BREAK.sub(r"\1\2", text)
    kept = [line for line in text.split("\n") if not _PAGE_MARKER.match(line)]
    text = _TRAILING_SPACE.sub("", "\n".join(kept))
    return _EXTRA_BLANKS.sub("\n\n", text).strip()


def _from_pdf(data: bytes) -> str:
    try:
        with fitz.open(stream=data, filetype="pdf") as document:
            if document.needs_pass:
                raise ExtractionError("the pdf is password protected")
            return "\n".join(page.get_text("text") for page in document)
    except ExtractionError:
        raise
    except Exception as exc:
        raise ExtractionError(f"could not read the pdf: {exc}") from exc


def _from_docx(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"could not read the docx: {exc}") from exc

    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            filled = [cell for cell in cells if cell]
            if filled:
                blocks.append(" | ".join(filled))
    return "\n".join(blocks)


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        raw = _from_pdf(data)
    elif name.endswith(".docx"):
        raw = _from_docx(data)
    elif name.endswith(".txt") or name.endswith(".md"):
        raw = data.decode("utf-8", errors="replace")
    else:
        raise ExtractionError("supported formats are pdf, docx, and txt")

    text = normalise(raw)
    if not text:
        raise ExtractionError(
            "no selectable text was found; the file is probably a scan"
        )
    return text
